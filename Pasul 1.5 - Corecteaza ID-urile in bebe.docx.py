# -*- coding: utf-8 -*-
"""
Pasul 1.5 - Corectează ID-urile în bebe.docx
Acest script caută fișierele RO corespunzătoare articolelor din bebe.docx
bazat pe titluri și actualizează ID-urile cu cele corecte.
"""

import os
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import unidecode

def normalize_title(title):
    """Normalizează titlul pentru comparare"""
    # Elimină diacriticele
    normalized = unidecode.unidecode(title.lower())
    # Elimină caracterele speciale
    normalized = re.sub(r'[^a-z0-9\s]', '', normalized)
    # Elimină spațiile multiple
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    return normalized

def generate_filename_from_title(title):
    """Generează numele fișierului din titlu (similar cu Pasul 2)"""
    normalized = unidecode.unidecode(title.lower())
    normalized = re.sub(r'[^a-z0-9\-]+', '-', normalized)
    normalized = re.sub(r'-+', '-', normalized).strip('-')
    return f"{normalized}.html"

def extract_id_from_html(file_path):
    """Extrage ID-ul din fișierul HTML"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        
        # Căutăm pattern-ul pentru ID
        patterns = [
            r'<!-- \$item_id = (\d+);',
            r'\$item_id = (\d+);'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, content)
            if match:
                return match.group(1)
        
        return None
    except Exception as e:
        print(f"Eroare la citirea {file_path}: {e}")
        return None

def extract_title_from_html(file_path):
    """Extrage titlul din fișierul HTML"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        
        soup = BeautifulSoup(content, 'html.parser')
        title_tag = soup.find('h1', class_='den_articol')
        if title_tag:
            return title_tag.get_text().strip()
        return None
    except Exception as e:
        print(f"Eroare la citirea titlului din {file_path}: {e}")
        return None

def build_ro_index(ro_dir):
    """Construiește un index al fișierelor RO cu titluri și ID-uri"""
    index = {}  # normalized_title -> (filename, id, original_title)
    filename_index = {}  # filename_base -> (filename, id, original_title)
    
    print(f"Indexare fișiere din: {ro_dir}")
    
    for filename in os.listdir(ro_dir):
        if not filename.endswith('.html'):
            continue
        
        file_path = os.path.join(ro_dir, filename)
        title = extract_title_from_html(file_path)
        item_id = extract_id_from_html(file_path)
        
        if title and item_id:
            normalized = normalize_title(title)
            index[normalized] = (filename, item_id, title)
            
            # Indexăm și după numele fișierului
            filename_base = os.path.splitext(filename)[0]
            filename_index[filename_base] = (filename, item_id, title)
    
    print(f"Indexate {len(index)} fișiere RO")
    return index, filename_index

def find_matching_ro_file(en_title, ro_index, ro_filename_index):
    """Găsește fișierul RO corespunzător titlului EN"""
    # Metoda 1: Normalizăm titlul EN și căutăm în index
    normalized_en = normalize_title(en_title)
    
    if normalized_en in ro_index:
        return ro_index[normalized_en]
    
    # Metoda 2: Generăm numele fișierului din titlu și căutăm
    expected_filename = generate_filename_from_title(en_title)
    expected_base = os.path.splitext(expected_filename)[0]
    
    if expected_base in ro_filename_index:
        return ro_filename_index[expected_base]
    
    # Metoda 3: Căutare parțială (cel puțin 3 cuvinte comune)
    en_words = set(normalized_en.split())
    best_match = None
    best_score = 0
    
    for norm_title, data in ro_index.items():
        ro_words = set(norm_title.split())
        common = len(en_words & ro_words)
        if common >= 3 and common > best_score:
            best_score = common
            best_match = data
    
    return best_match

def correct_ids_in_docx(docx_path, ro_dir, output_path=None):
    """Corectează ID-urile în documentul Word"""
    if output_path is None:
        output_path = docx_path  # Suprascrie fișierul original
    
    # Construim indexul fișierelor RO
    ro_index, ro_filename_index = build_ro_index(ro_dir)
    
    # Deschidem documentul
    doc = Document(docx_path)
    
    corrections = []
    current_title = None
    current_title_para_idx = None
    
    print(f"\nProcesare {docx_path}...")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Detectăm titlul (centrat, nu începe cu ID:)
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and text and not text.startswith("ID:"):
            current_title = text
            current_title_para_idx = i
        
        # Detectăm linia de ID
        elif text.startswith("ID:") and current_title:
            old_id = text.replace("ID:", "").strip()
            
            # Căutăm fișierul RO corespunzător
            match = find_matching_ro_file(current_title, ro_index, ro_filename_index)
            
            if match:
                filename, correct_id, ro_title = match
                
                if old_id != correct_id:
                    print(f"\n[CORECȚIE] {current_title[:50]}...")
                    print(f"  ID vechi: {old_id} -> ID corect: {correct_id}")
                    print(f"  Fișier RO: {filename}")
                    
                    # Actualizăm textul paragrafului
                    para.clear()
                    run = para.add_run(f"ID: {correct_id}")
                    run.font.size = para.runs[0].font.size if para.runs else None
                    
                    corrections.append({
                        'title': current_title,
                        'old_id': old_id,
                        'new_id': correct_id,
                        'ro_file': filename
                    })
                else:
                    print(f"[OK] {current_title[:50]}... -> ID: {correct_id}")
            else:
                print(f"[ATENȚIE] Nu s-a găsit fișier RO pentru: {current_title[:50]}...")
            
            current_title = None
    
    # Salvăm documentul
    doc.save(output_path)
    
    print(f"\n{'='*60}")
    print(f"Corecții efectuate: {len(corrections)}")
    print(f"Document salvat: {output_path}")
    
    return corrections

def main():
    # Căi
    bebe_docx = "bebe.docx"
    ro_dir = r"e:\Carte\BB\17 - Site Leadership\Principal\ro"
    
    if not os.path.exists(bebe_docx):
        print(f"Eroare: Fișierul {bebe_docx} nu există!")
        return
    
    if not os.path.exists(ro_dir):
        print(f"Eroare: Directorul {ro_dir} nu există!")
        return
    
    print("="*60)
    print("PASUL 1.5 - Corectare ID-uri în bebe.docx")
    print("="*60)
    
    # Corectăm ID-urile
    corrections = correct_ids_in_docx(bebe_docx, ro_dir)
    
    if corrections:
        print("\n" + "="*60)
        print("REZUMAT CORECȚII:")
        print("="*60)
        for c in corrections:
            print(f"  {c['title'][:40]}... : {c['old_id']} -> {c['new_id']}")
    
    print("\nProcesare completă!")
    print("Acum poți rula Pasul 2 pentru a regenera fișierele HTML cu ID-uri corecte.")

if __name__ == "__main__":
    main()
