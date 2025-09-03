import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from tqdm import tqdm
import time

# Lista fișierelor specifice de procesat  # toate articolele sunt preluate din  e:\Carte\BB\17 - Site Leadership\Principal\ro\
SPECIFIC_FILES = [
    'probatio-suprema.html',
    'legea-sigiliului-sacru.html',
    'iuramentum-bellatoris-et-ardor-cordis.html',
    'aseisthos.html',
    'principiul-fons-serenitatis.html',
    'opus-spiritualis.html',
    'visio-aurea.html',
    'aedaud.html',
    'sanctuarium-silentii.html',
    'narratio-divinae.html',
    'sanguis-litterae.html',
    'povestea-devine-sangele-celui-ce-scrie.html',
    'doar-uita-te-si-vezi.html',
    'scriptorem-serpentis.html',
    'scriptura-arcanum.html',
    'magister-verbi.html',
    'testamentum-ignis.html',
    'fabula-perpetua.html',
    'semna-divinae-puritatis.html'
]

def process_paragraph(paragraph, p_tag, is_bold=False):
    """Procesează un paragraf și aplică formatarea corectă"""

    # Determină dacă întregul paragraf trebuie să fie bold bazat pe clasa sa
    is_text_obisnuit2 = 'text_obisnuit2' in p_tag.get('class', [])

    # Procesează fiecare element din paragraf
    for element in p_tag.children:
        if isinstance(element, str):
            # Text simplu
            run = paragraph.add_run(element)
            run.bold = is_text_obisnuit2
        elif element.name == 'em':
            # Text italic
            run = paragraph.add_run(element.get_text())
            run.italic = True
            run.bold = is_text_obisnuit2
        elif element.name == 'span' and 'text_obisnuit2' in element.get('class', []):
            # Text în span cu clasa text_obisnuit2 - trebuie să fie bold
            run = paragraph.add_run(element.get_text())
            run.bold = True
        else:
            # Alte elemente - procesare normală
            run = paragraph.add_run(element.get_text() if element.name else str(element))
            run.bold = is_text_obisnuit2

def extract_item_id(content):
    """Extrage ID-ul articolului din comentariul HTML"""
    id_pattern = re.compile(r'<!-- \$item_id = (\d+); // Replace that with your rating id -->')
    match = id_pattern.search(content)
    if match:
        return match.group(1)
    return "N/A"  # Returnează N/A dacă nu găsește ID-ul

def process_html_file(file_path, document):
    print(f"\nProcesare fișier: {os.path.basename(file_path)}")
    start_time = time.time()

    # Încearcă mai multe codificări
    codecs_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    content = None

    for codec in codecs_to_try:
        try:
            with open(file_path, 'r', encoding=codec, errors='replace') as file:
                content = file.read()
                print(f"Fișier citit cu succes folosind codificarea: {codec}")
                break  # Am găsit o codificare care funcționează
        except UnicodeDecodeError:
            print(f"Nu s-a putut citi fișierul cu codificarea: {codec}")
            continue

    if content is None:
        print(f"EROARE: Nu s-a putut citi fișierul cu niciuna dintre codificările încercate.")
        document.add_paragraph(f"EROARE LA PROCESARE: {os.path.basename(file_path)}")
        return

    # Extrage ID-ul articolului
    item_id = extract_item_id(content)
    print(f"ID articol: {item_id}")

    soup = BeautifulSoup(content, 'html.parser')

    # Procesează titlul
    title = soup.find('h1', class_='den_articol')
    if title:
        title_text = title.text
        print(f"Adăugare titlu: {title_text[:50]}...")

        # Crează un paragraf pentru titlu
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adaugă titlul cu formatare
        run = title_paragraph.add_run(title_text)
        run.bold = True
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0)

        # Adaugă ID-ul articolului doar sub titlu
        id_paragraph = document.add_paragraph()
        id_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        id_run = id_paragraph.add_run(f"ID: {item_id}")
        id_run.font.color.rgb = RGBColor(128, 128, 128)  # Culoare gri
        id_run.font.size = Pt(8)  # Dimensiune mică pentru text

    # Găsește conținutul între markeri
    start_marker = '<!-- ARTICOL START -->'
    end_marker = '<!-- ARTICOL FINAL -->'
    start_idx = content.find(start_marker)
    end_idx = content.find(end_marker)

    paragraphs_processed = 0
    if start_idx != -1 and end_idx != -1:
        article_content = content[start_idx + len(start_marker):end_idx]
        article_soup = BeautifulSoup(article_content, 'html.parser')

        total_paragraphs = len(article_soup.find_all('p'))
        print(f"Găsite {total_paragraphs} paragrafe pentru procesare")

        for p in article_soup.find_all('p'):
            try:
                paragraphs_processed += 1
                paragraph = document.add_paragraph()
                process_paragraph(paragraph, p)
            except Exception as e:
                print(f"Eroare la procesarea paragrafului {paragraphs_processed}: {e}")
                paragraph = document.add_paragraph("EROARE LA PROCESAREA PARAGRAFULUI")
    else:
        print("Nu s-au găsit markerii pentru conținutul articolului")
        document.add_paragraph("EROARE: Nu s-au găsit markerii pentru conținutul articolului")

    document.add_paragraph()

    end_time = time.time()
    print(f"Procesate {paragraphs_processed} paragrafe în {end_time - start_time:.2f} secunde")

def main():
    input_folder = r'e:\Carte\BB\17 - Site Leadership\Principal\ro'
    output_file = 'articole_compilate.docx'

    print("\nÎncepere procesare articole HTML specificate...")
    print("=" * 50)

    # Verifică existența fișierelor
    files_to_process = []
    missing_files = []

    for filename in SPECIFIC_FILES:
        file_path = os.path.join(input_folder, filename)
        if os.path.exists(file_path):
            files_to_process.append(file_path)
        else:
            missing_files.append(filename)

    print(f"\nGăsite {len(files_to_process)} din {len(SPECIFIC_FILES)} fișiere specificate")

    if missing_files:
        print("\nATENȚIE! Următoarele fișiere nu au fost găsite:")
        for file in missing_files:
            print(f"- {file}")

    document = Document()

    # Crează progress bar
    with tqdm(total=len(files_to_process), desc="Progres total") as pbar:
        for file_path in files_to_process:
            try:
                process_html_file(file_path, document)
            except Exception as e:
                print(f"\nEroare la procesarea fișierului {os.path.basename(file_path)}: {e}")
                document.add_paragraph(f"EROARE LA PROCESARE: {os.path.basename(file_path)}")
                document.add_paragraph(f"Detalii eroare: {str(e)}")
            pbar.update(1)

    print("\nSalvare document final...")
    document.save(output_file)
    print(f"\nDocument creat cu succes: {output_file}")
    print(f"Au fost procesate {len(files_to_process)} fișiere")
    print("\nProcesare completă!")

if __name__ == "__main__":
    main()